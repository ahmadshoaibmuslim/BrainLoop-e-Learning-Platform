from docx import Document
import re

md_path = 'BrainLoop_FYP_Poster_Content.md'
docx_path = 'BrainLoop_FYP_Poster_Content.docx'

text = open(md_path, encoding='utf-8').read()
text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
text = re.sub(r'\*([^*]+)\*', r'\1', text)
text = re.sub(r'`([^`]+)`', r'\1', text)

lines = text.splitlines()
doc = Document()
current_table = None

for line in lines:
    if line.startswith('# '):
        current_table = None
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        current_table = None
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        current_table = None
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith('#### '):
        current_table = None
        doc.add_heading(line[5:].strip(), level=4)
    elif line.startswith('- '):
        current_table = None
        doc.add_paragraph(line[2:].strip(), style='List Bullet')
    elif line.startswith('* '):
        current_table = None
        doc.add_paragraph(line[2:].strip(), style='List Bullet')
    elif line.startswith('|') and '---' not in line:
        cells = [c.strip() for c in line.strip('|').split('|')]
        if current_table is None:
            current_table = doc.add_table(rows=0, cols=len(cells))
        row = current_table.add_row().cells
        for i, cell in enumerate(cells):
            row[i].text = cell
    elif line.strip() == '':
        current_table = None
        doc.add_paragraph('')
    elif line.strip().startswith('```'):
        continue
    else:
        if line.strip().startswith('> '):
            p = doc.add_paragraph(line.strip()[2:])
            p.style = 'Intense Quote'
        else:
            doc.add_paragraph(line.strip())


doc.save(docx_path)
print(f'Saved {docx_path}')
